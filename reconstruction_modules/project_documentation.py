# ======================================================
# project_documentation.py - Professional Project Documentation Generator
# ======================================================

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_project_documentation(output_path=None):
    """Generate professional Word document for project documentation"""
    
    if output_path is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f"Gaza_Reconstruction_DSS_Documentation_{timestamp}.docx"
    
    doc = Document()
    
    # ===== COVER PAGE =====
    # Title
    title = doc.add_heading('Gaza Strip Reconstruction', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_run.bold = True
    
    # Subtitle
    subtitle = doc.add_heading('AI-Enhanced Decision Support System', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(139, 0, 0)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project Info Box
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(
        '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
        'Advanced Geospatial Analysis & Prioritization Framework\n'
        'For Post-Conflict Reconstruction Planning\n'
        '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
    )
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(70, 70, 70)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Version and Date
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.add_run(f'Version 1.0\n{datetime.now().strftime("%B %Y")}')
    version_run.font.size = Pt(14)
    version_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()
    
    # ===== EXECUTIVE SUMMARY =====
    doc.add_heading('Executive Summary', 1)
    
    summary_text = """
The Gaza Strip Reconstruction Decision Support System (DSS) is an advanced, AI-enhanced geospatial 
analysis platform designed to prioritize and optimize post-conflict reconstruction efforts. This 
system integrates multiple data sources, applies sophisticated algorithms, and generates actionable 
insights to support evidence-based decision-making for reconstruction planning.

The system addresses the critical challenge of efficiently allocating limited resources across 
extensive damage areas by providing:
• Objective, data-driven priority rankings for reconstruction zones
• Comprehensive infrastructure assessment and mapping
• Multi-criteria decision analysis incorporating damage severity, population impact, and strategic importance
• Phased implementation plans with cost estimates and timelines
• Interactive visualizations and detailed analytical reports
"""
    
    doc.add_paragraph(summary_text.strip())
    doc.add_paragraph()
    
    # ===== KEY FEATURES =====
    doc.add_heading('Key Features', 2)
    
    features = [
        ('AI-Enhanced Scoring Algorithm', 'Machine learning-based prioritization using weighted multi-criteria analysis'),
        ('Geospatial Analysis', 'Hexagonal grid-based spatial partitioning for uniform zone assessment'),
        ('Infrastructure Mapping', 'Comprehensive mapping of hospitals, schools, utilities, and transportation networks'),
        ('Damage Assessment', 'Temporal damage tracking with age-based ranking algorithms'),
        ('Strategic Planning', 'Municipality-specific reconstruction strategies based on damage patterns'),
        ('Interactive Visualizations', 'Dynamic maps with multiple layers and real-time data exploration'),
        ('Phased Implementation', 'Four-phase reconstruction timeline (Emergency, Basics, Development, Improvement)'),
        ('Cost Estimation', 'Automated cost calculation based on infrastructure type and damage severity')
    ]
    
    for feature, description in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{feature}: ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # ===== DATA SOURCES =====
    doc.add_heading('Data Sources & Methodology', 1)
    
    doc.add_heading('Primary Data Sources', 2)
    
    data_sources = [
        {
            'name': 'UNOSAT Damage Assessment',
            'source': 'United Nations Satellite Centre (UNOSAT)',
            'description': 'Comprehensive damage site mapping with temporal tracking across multiple assessment cycles',
            'file': 'UNOSAT_GazaStrip_CDA_11October2025.gdb',
            'records': '~50,000+ damage sites',
            'coverage': 'Gaza Strip (October 2023 - October 2025)'
        },
        {
            'name': 'UNOSAT Road Damage Assessment',
            'source': 'UNOSAT RoadCDA',
            'description': 'Detailed street network damage assessment with severity classifications',
            'file': 'CE20231007PSE_UNOSAT_GazaStrip_RoadCDA_20250708_GDB_v1.gdb',
            'records': '61,892 street segments',
            'coverage': 'Complete road network with damage levels (0-4 scale)'
        },
        {
            'name': 'Infrastructure Facilities',
            'source': 'Humanitarian Data Exchange (HDX)',
            'description': 'Critical infrastructure locations including healthcare, education, and utilities',
            'file': 'Multiple GeoPackage files',
            'records': '8 major hospitals, 6 universities, 100+ schools, 12 water utilities, 6 power facilities',
            'coverage': 'Verified facility locations with operational status'
        },
        {
            'name': 'Population Data',
            'source': 'UN OCHA + Estimates',
            'description': 'Population density, displacement rates, and IDP concentration areas',
            'file': 'Integrated from multiple humanitarian sources',
            'records': 'Zone-level population estimates',
            'coverage': 'Pre-war baseline and current population distribution'
        }
    ]
    
    for ds in data_sources:
        doc.add_heading(ds['name'], 3)
        p = doc.add_paragraph()
        p.add_run('Source: ').bold = True
        p.add_run(f"{ds['source']}\n")
        p.add_run('Description: ').bold = True
        p.add_run(f"{ds['description']}\n")
        p.add_run('Data File: ').bold = True
        p.add_run(f"{ds['file']}\n")
        p.add_run('Records: ').bold = True
        p.add_run(f"{ds['records']}\n")
        p.add_run('Coverage: ').bold = True
        p.add_run(ds['coverage'])
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ===== TECHNICAL ARCHITECTURE =====
    doc.add_heading('Technical Architecture', 1)
    
    doc.add_heading('System Components', 2)
    
    components = [
        ('Data Loading & Preprocessing', 'data_loader.py', 'Loads geospatial data, filters Gaza boundaries, corrects municipality names'),
        ('Spatial Analysis Engine', 'spatial_analysis.py', 'Creates hexagonal grid (500m radius), integrates damage and infrastructure data'),
        ('AI Scoring Engine', 'scoring_engine.py', 'Calculates priority scores using weighted multi-criteria analysis'),
        ('Municipality Analysis', 'municipality_analysis.py', 'Determines reconstruction strategies based on damage patterns'),
        ('Project Generator', 'project_generator.py', 'Generates reconstruction projects with cost and timeline estimates'),
        ('Visualization Engine', 'visualization.py', 'Creates interactive maps with multiple infrastructure layers'),
        ('Phased Implementation', 'phased_implementation.py', 'Divides projects into 4 phases over 84 months'),
        ('Reporting System', 'Multiple report generators', 'Produces PDF reports and Excel exports')
    ]
    
    for component, module, description in components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{component} ').bold = True
        p.add_run(f'({module})\n')
        p.add_run(description)
    
    doc.add_paragraph()
    
    doc.add_heading('Core Algorithms', 2)
    
    algorithms = [
        ('Hexagonal Grid Partitioning', 'Divides Gaza Strip into uniform 500m hexagonal zones for consistent analysis'),
        ('AI Priority Scoring', 'Weighted formula: AI_Score = Σ(Feature_i × Weight_i) with 12 weighted factors'),
        ('Damage Age Ranking', 'Temporal analysis of damage persistence and severity trends'),
        ('Street Priority Algorithm', 'Priority = road_type×3 + damage×2 + major_artery×10 + strategic_score×4'),
        ('Municipality Correction', 'Coordinate-based boundary matching for accurate location assignment'),
        ('Multi-Phase Planning', 'Divides projects into Emergency (0-12mo), Basics (12-30mo), Development (30-54mo), Improvement (54-84mo)')
    ]
    
    for algo, description in algorithms:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{algo}: ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # ===== SCORING METHODOLOGY =====
    doc.add_heading('AI Scoring Methodology', 1)
    
    doc.add_paragraph(
        'The system employs a sophisticated multi-criteria decision analysis (MCDA) approach '
        'to calculate reconstruction priorities. Each zone receives an AI Score based on weighted factors:'
    )
    doc.add_paragraph()
    
    # Scoring weights table
    doc.add_heading('Scoring Weights', 2)
    
    weights = [
        ('Damage Severity', '20%', 'Number and intensity of damage sites'),
        ('Major Hospitals', '40% × 1.7', 'Proximity to critical healthcare facilities'),
        ('Healthcare Clinics', '10% × 0.3', 'Access to medical services'),
        ('Water Infrastructure', '15%', 'Water treatment and distribution facilities'),
        ('Power Infrastructure', '12%', 'Electricity generation and distribution'),
        ('Street Network', '10%', 'Transportation and access routes'),
        ('Schools', '6%', 'Primary and secondary education facilities'),
        ('Universities', '8%', 'Higher education institutions'),
        ('Municipal Services', '4%', 'Government and administrative facilities'),
        ('Population Density', '15%', 'Current population concentration'),
        ('Education Centers', '2%', 'Additional educational facilities'),
        ('Damage Age', '25%', 'Temporal persistence of damage')
    ]
    
    for criterion, weight, description in weights:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{criterion} ({weight}): ').bold = True
        p.add_run(description)
    
    doc.add_paragraph()
    
    doc.add_heading('Priority Classification', 2)
    
    doc.add_paragraph('Zones are classified into six priority levels:')
    
    priorities = [
        ('🔴 Critical', 'Top 10%', 'Immediate reconstruction required'),
        ('🔴 Very High', '10-20%', 'Urgent reconstruction needed'),
        ('🟠 High', '20-40%', 'High priority reconstruction'),
        ('🟡 Medium', '40-60%', 'Medium priority reconstruction'),
        ('🟢 Low', '60-80%', 'Lower priority reconstruction'),
        ('🟢 Very Low', '80-100%', 'Lowest priority reconstruction')
    ]
    
    for level, range_pct, description in priorities:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{level} ({range_pct}): ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # ===== OUTPUTS =====
    doc.add_heading('System Outputs', 1)
    
    doc.add_heading('Interactive Maps', 2)
    
    maps = [
        ('Reconstruction Priority Map', 'Color-coded zones with priority rankings, complete street network, infrastructure layers'),
        ('Damage Heatmap', 'Gradient visualization of damage intensity across Gaza Strip'),
        ('Streets Damage Assessment', 'Detailed street-level damage classification (Severe/High/Medium/Low)'),
        ('Streets Reconstruction Plan', 'Priority-based street reconstruction with cost and timeline estimates'),
        ('Phased Implementation Map', 'Projects grouped by implementation phase with timeline visualization'),
        ('3D Damage Visualization', 'Three-dimensional representation of damage distribution')
    ]
    
    for map_name, description in maps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{map_name}: ').bold = True
        p.add_run(description)
    
    doc.add_paragraph()
    
    doc.add_heading('Reports & Data Exports', 2)
    
    outputs = [
        ('All Maps Summary Report (PDF)', 'Comprehensive analysis of all visualizations with statistics'),
        ('Reconstruction Projects (Excel)', 'Detailed project list with costs, timelines, and priorities'),
        ('Streets Reconstruction (Excel)', 'Street-specific reconstruction plan with 100 top-priority streets'),
        ('Phased Projects (Excel)', 'Projects organized by implementation phase'),
        ('Phased Timeline Chart (PNG)', 'Visual timeline showing project distribution across 4 phases'),
        ('Statistical Dashboard (PNG)', 'Key metrics and distribution charts')
    ]
    
    for output, description in outputs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{output}: ').bold = True
        p.add_run(description)
    
    doc.add_page_break()
    
    # ===== TECHNICAL SPECIFICATIONS =====
    doc.add_heading('Technical Specifications', 1)
    
    doc.add_heading('Software Stack', 2)
    
    tech_stack = [
        ('Programming Language', 'Python 3.8+'),
        ('Geospatial Processing', 'GeoPandas, Shapely, Fiona'),
        ('Data Analysis', 'Pandas, NumPy, Scikit-learn'),
        ('Visualization', 'Folium, Matplotlib, Seaborn, Plotly'),
        ('Reporting', 'ReportLab (PDF), OpenPyXL (Excel), python-docx (Word)'),
        ('Coordinate Systems', 'EPSG:32636 (UTM Zone 36N), EPSG:4326 (WGS84)')
    ]
    
    for component, technology in tech_stack:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{component}: ').bold = True
        p.add_run(technology)
    
    doc.add_paragraph()
    
    doc.add_heading('System Requirements', 2)
    
    requirements = [
        'Python 3.8 or higher',
        'Minimum 8GB RAM (16GB recommended)',
        'Minimum 5GB free disk space',
        'Internet connection for basemap tiles',
        'Modern web browser for interactive maps'
    ]
    
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== USAGE GUIDE =====
    doc.add_heading('Usage Guide', 1)
    
    doc.add_heading('Running the System', 2)
    
    doc.add_paragraph('1. Ensure all data files are in the correct directories:')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('UNOSAT_GazaStrip_CDA_11October2025.gdb').italic = True
    p.add_run(' (main directory)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('CE20231007PSE_UNOSAT_GazaStrip_RoadCDA_20250708_GDB_v1.gdb').italic = True
    p.add_run(' (reconstruction_modules directory)')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Infrastructure files').italic = True
    p.add_run(' (reconstruction_modules/gaza_gis_data directory)')
    
    doc.add_paragraph()
    doc.add_paragraph('2. Run the main script:')
    code = doc.add_paragraph('python reconstruction_modules/main.py')
    code.style = 'Intense Quote'
    
    doc.add_paragraph()
    doc.add_paragraph('3. Follow the interactive menu to select project types')
    doc.add_paragraph('4. Review generated outputs in the main directory')
    
    doc.add_paragraph()
    
    doc.add_heading('Interpreting Results', 2)
    
    doc.add_paragraph(
        'The system generates prioritized reconstruction zones with AI scores ranging from 0 to 1. '
        'Higher scores indicate higher reconstruction priority based on multiple factors including '
        'damage severity, infrastructure density, population impact, and strategic importance.'
    )
    
    doc.add_page_break()
    
    # ===== CONCLUSION =====
    doc.add_heading('Conclusion', 1)
    
    conclusion = """
The Gaza Strip Reconstruction Decision Support System represents a comprehensive, data-driven approach 
to post-conflict reconstruction planning. By integrating advanced geospatial analysis, artificial 
intelligence algorithms, and multi-criteria decision analysis, the system provides objective, 
evidence-based prioritization for reconstruction efforts.

The system's key strengths include:
• Objective, transparent prioritization methodology
• Integration of multiple authoritative data sources
• Comprehensive infrastructure assessment
• Flexible, phased implementation planning
• Interactive visualizations for stakeholder engagement
• Detailed cost and timeline estimates

This tool is designed to support decision-makers, planners, donors, and implementing agencies in 
efficiently allocating resources and coordinating reconstruction efforts to maximize impact and 
accelerate recovery.
"""
    
    doc.add_paragraph(conclusion.strip())
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n'
        'Gaza Strip Reconstruction Decision Support System\n'
        f'Generated: {datetime.now().strftime("%B %d, %Y")}\n'
        '━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━'
    )
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Save document
    doc.save(output_path)
    print(f"   - Project documentation saved to {output_path}")
    return output_path
